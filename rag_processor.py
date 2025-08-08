import os
import requests
import gc
from typing import List
import email
from email.parser import BytesParser
from email.policy import default
import io
import numpy as np

# Lightweight document processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# LangChain essentials
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.documents import Document

# Vector storage - try FAISS first, fallback to simple search
try:
    from sentence_transformers import SentenceTransformer
    import faiss
    FAISS_AVAILABLE = True
except ImportError:
    FAISS_AVAILABLE = False


class SimpleEmbeddings:
    """Lightweight embedding class"""
    def __init__(self):
        if FAISS_AVAILABLE:
            self.model = SentenceTransformer('all-MiniLM-L6-v2')
        else:
            self.model = None
    
    def embed_documents(self, texts):
        if self.model:
            return self.model.encode(texts, convert_to_tensor=False).tolist()
        else:
            # Fallback: return dummy embeddings
            return [[0.0] * 384 for _ in texts]
    
    def embed_query(self, text):
        if self.model:
            return self.model.encode([text], convert_to_tensor=False)[0].tolist()
        else:
            return [0.0] * 384


class RAGProcessor:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.llm = None  # Lazy load
        self.embeddings = None  # Lazy load
        
        # Simple RAG prompt template
        self.prompt_template = """You are an assistant for question-answering tasks. Use the following pieces of retrieved context to answer the question. If you don't know the answer, just say that you don't know. Keep the answer concise and accurate.

Context:
{context}

Question: {question}

Answer:"""

    def _get_llm(self):
        """Lazy load LLM to save memory"""
        if self.llm is None:
            self.llm = ChatGoogleGenerativeAI(
                model="gemini-1.5-flash",  # Faster and cheaper than gemini-pro
                google_api_key=self.api_key,
                temperature=0.1
            )
        return self.llm

    def _get_embeddings(self):
        """Lazy load embeddings with lightweight model"""
        if self.embeddings is None:
            self.embeddings = SimpleEmbeddings()
        return self.embeddings

    def _detect_file_type(self, url: str) -> str:
        """Detect file type from URL"""
        url_lower = url.lower()
        if '.pdf' in url_lower or 'pdf' in url_lower:
            return 'pdf'
        elif '.docx' in url_lower or 'word' in url_lower:
            return 'docx'
        elif '.eml' in url_lower or 'email' in url_lower:
            return 'eml'
        
        # Try to check content-type header
        try:
            response = requests.head(url, timeout=5)
            content_type = response.headers.get('content-type', '').lower()
            if 'pdf' in content_type:
                return 'pdf'
            elif 'word' in content_type or 'docx' in content_type:
                return 'docx'
        except:
            pass
        
        return 'pdf'  # Default assumption

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF bytes"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 not available for PDF processing")
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available for DOCX processing")
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")

    def _extract_text_from_eml(self, file_content: bytes) -> str:
        """Extract text from EML bytes"""
        try:
            msg = BytesParser(policy=default).parsebytes(file_content)
            text = ""
            
            # Get subject
            if msg['Subject']:
                text += f"Subject: {msg['Subject']}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            text += payload.decode('utf-8', errors='ignore')
                    elif part.get_content_type() == "text/html" and BS4_AVAILABLE:
                        payload = part.get_payload(decode=True)
                        if payload:
                            soup = BeautifulSoup(payload, 'html.parser')
                            text += soup.get_text()
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    text += payload.decode('utf-8', errors='ignore')
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting EML text: {str(e)}")

    def _download_and_extract_text(self, url: str) -> str:
        """Download file from URL and extract text"""
        file_type = self._detect_file_type(url)
        
        # Download file content
        response = requests.get(url, timeout=30)
        if response.status_code != 200:
            raise Exception(f"Failed to download file: {response.status_code}")
        
        file_content = response.content
        
        # Extract text based on file type
        if file_type == 'pdf':
            return self._extract_text_from_pdf(file_content)
        elif file_type == 'docx':
            return self._extract_text_from_docx(file_content)
        elif file_type == 'eml':
            return self._extract_text_from_eml(file_content)
        else:
            # Try PDF as fallback
            try:
                return self._extract_text_from_pdf(file_content)
            except:
                # If all fails, try to decode as text
                return file_content.decode('utf-8', errors='ignore')

    def process_document_from_url(self, doc_url: str, questions: List[str]) -> List[str]:
        try:
            # === Step 1: Download and extract text ===
            document_text = self._download_and_extract_text(doc_url)
            
            if not document_text.strip():
                raise Exception("No text could be extracted from the document")

            # === Step 2: Split into chunks ===
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=400,
                chunk_overlap=50,
                length_function=len
            )
            
            # Create document objects
            docs = [Document(page_content=document_text)]
            chunks = splitter.split_documents(docs)
            
            if not chunks:
                raise Exception("Document could not be split into chunks")

            # Clear original docs to free memory
            del docs, document_text
            gc.collect()

            # === Step 3: Create vector store ===
            embeddings = self._get_embeddings()
            
            # Extract texts for embedding
            texts = [chunk.page_content for chunk in chunks]
            
            if FAISS_AVAILABLE:
                # Create FAISS index
                embeddings_vectors = embeddings.embed_documents(texts)
                dimension = len(embeddings_vectors[0])
                index = faiss.IndexFlatL2(dimension)
                index.add(np.array(embeddings_vectors, dtype=np.float32))
                
                # Store texts with index
                self.chunk_texts = texts
                self.faiss_index = index
                
            else:
                # Fallback - just store chunks for simple retrieval
                self.chunk_texts = texts

            # Clear chunks to free memory
            del chunks, texts
            gc.collect()

            # === Step 4: Process questions ===
            llm = self._get_llm()
            parser = StrOutputParser()
            
            answers = []
            for question in questions:
                try:
                    # Get relevant context
                    if FAISS_AVAILABLE:
                        context = self._get_relevant_context_faiss(question, k=3)
                    else:
                        context = self._get_relevant_context_simple(question, k=3)
                    
                    # Create prompt
                    prompt = self.prompt_template.format(context=context, question=question)
                    
                    # Get answer
                    response = llm.invoke(prompt)
                    answer = parser.invoke(response)
                    answers.append(answer.strip())
                    
                    # Force garbage collection between questions
                    gc.collect()
                    
                except Exception as e:
                    answers.append(f"Error processing question: {str(e)}")

            return answers

        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")

    def _get_relevant_context_faiss(self, question: str, k: int = 3) -> str:
        """Get relevant context using FAISS"""
        try:
            # Embed the question
            question_embedding = self.embeddings.embed_query(question)
            
            # Search FAISS index
            distances, indices = self.faiss_index.search(
                np.array([question_embedding], dtype=np.float32), k
            )
            
            # Get relevant texts
            relevant_texts = [self.chunk_texts[idx] for idx in indices[0] if idx < len(self.chunk_texts)]
            return "\n\n".join(relevant_texts)
            
        except Exception:
            # Fallback to first few chunks
            return "\n\n".join(self.chunk_texts[:k])

    def _get_relevant_context_simple(self, question: str, k: int = 3) -> str:
        """Simple keyword-based context retrieval"""
        try:
            # Simple keyword matching
            question_lower = question.lower()
            question_words = set(question_lower.split())
            
            # Score chunks based on keyword overlap
            scored_chunks = []
            for i, chunk in enumerate(self.chunk_texts):
                chunk_lower = chunk.lower()
                chunk_words = set(chunk_lower.split())
                overlap = len(question_words.intersection(chunk_words))
                scored_chunks.append((overlap, i, chunk))
            
            # Sort by score and take top k
            scored_chunks.sort(reverse=True, key=lambda x: x[0])
            relevant_chunks = [chunk for _, _, chunk in scored_chunks[:k]]
            
            if not relevant_chunks:
                relevant_chunks = self.chunk_texts[:k]
            
            return "\n\n".join(relevant_chunks)
            
        except Exception:
            # Ultimate fallback
            return "\n\n".join(self.chunk_texts[:k])